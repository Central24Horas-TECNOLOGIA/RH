from __future__ import annotations

import io
import json
import logging
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from zipfile import BadZipFile

from .helpers import normalize_compare_text, normalize_text


logger = logging.getLogger(__name__)

CV_MAX_BYTES = 10 * 1024 * 1024
SUPPORTED_CV_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
}
IMAGE_CV_EXTENSIONS = set()
DOC_CONVERTER_FAILURE_MESSAGE = (
    "Não foi possível converter o arquivo .doc neste servidor. Verifique se o Microsoft Word ou LibreOffice "
    "está instalado/configurado, ou envie o currículo em PDF ou DOCX."
)


class CvTextExtractionError(Exception):
    def __init__(self, user_message: str, *, technical_message: str = ""):
        super().__init__(technical_message or user_message)
        self.user_message = user_message
        self.technical_message = technical_message or user_message


CV_KEYWORDS = [
    "excel",
    "word",
    "power bi",
    "powerbi",
    "atendimento",
    "suporte",
    "planilha",
    "dashboards",
    "dashboard",
    "sql",
    "python",
    "javascript",
    "react",
    "talkdesk",
    "power automate",
    "sharepoint",
    "teams",
    "analise",
    "dados",
    "cliente",
    "monitoramento",
    "organizacao",
    "comunicacao",
    "responsabilidade",
    "proatividade",
    "pontualidade",
    "trabalho em equipe",
    "empatia",
    "cordialidade",
    "resiliencia",
    "agilidade",
    "atencao aos detalhes",
    "crm",
    "help desk",
    "control desk",
    "atendimento telefonico",
    "atendimento digital",
    "rotinas administrativas",
    "digitacao",
    "relatorios",
    "escala",
    "call center",
    "operacao",
]

CV_WEIGHTS_BY_ROLE = {
    "Control Desk": {
        "keywords": {
            "control desk": 3.0,
            "monitoramento": 2.0,
            "atendimento": 2.0,
            "atendimento telefonico": 1.5,
            "atendimento digital": 1.5,
            "suporte": 1.5,
            "help desk": 1.5,
            "crm": 1.2,
            "excel": 1.2,
            "relatorios": 1.0,
            "escala": 1.0,
            "operacao": 1.0,
            "call center": 1.0,
        },
        "min_keywords": 3,
    },
    "Jovem Aprendiz": {
        "keywords": {
            "excel": 1.5,
            "word": 1.0,
            "atendimento": 2.0,
            "cliente": 1.5,
            "dados": 1.0,
        },
        "min_keywords": 2,
    },
    "Operador": {
        "keywords": {
            "atendimento": 2.5,
            "cliente": 2.0,
            "suporte": 1.5,
            "talkdesk": 2.0,
            "monitoramento": 1.0,
            "excel": 1.0,
        },
        "min_keywords": 2,
    },
    "Supervisor": {
        "keywords": {
            "excel": 2.0,
            "atendimento": 2.0,
            "cliente": 1.5,
            "monitoramento": 1.5,
            "dados": 1.5,
            "power bi": 2.0,
        },
        "min_keywords": 3,
    },
    "Estagiario": {
        "keywords": {
            "excel": 1.5,
            "python": 2.0,
            "sql": 2.0,
            "dados": 1.5,
            "javascript": 1.5,
            "react": 1.5,
        },
        "min_keywords": 2,
    },
    "Analista": {
        "keywords": {
            "excel": 1.5,
            "power bi": 2.5,
            "sql": 2.0,
            "python": 2.0,
            "dados": 2.0,
            "analise": 1.5,
        },
        "min_keywords": 3,
    },
    "TI": {
        "keywords": {
            "python": 2.5,
            "javascript": 2.0,
            "react": 2.0,
            "sql": 2.0,
            "sharepoint": 1.5,
            "power automate": 1.5,
            "teams": 1.0,
            "suporte": 1.5,
        },
        "min_keywords": 3,
    },
}

EMAIL_REGEX = re.compile(r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})", re.IGNORECASE)
PHONE_REGEX = re.compile(r"(?:(?:\+?55)\s*)?(?:\(?\d{2}\)?\s*)?(?:9?\d{4})[-\s.]?\d{4}")
WHATSAPP_REGEX = re.compile(r"(?:(?:\+?55)\s*)?(?:\(?\d{2}\)?\s*)?(?:9\d{4})[-\s.]?\d{4}")

BEHAVIORAL_COMPETENCIES = [
    "comunicacao",
    "organizacao",
    "responsabilidade",
    "proatividade",
    "pontualidade",
    "trabalho em equipe",
    "atendimento ao cliente",
    "empatia",
    "cordialidade",
    "resiliencia",
    "agilidade",
    "atencao aos detalhes",
]

TECHNICAL_COMPETENCIES = [
    "excel",
    "word",
    "sistemas internos",
    "atendimento telefonico",
    "atendimento digital",
    "crm",
    "help desk",
    "control desk",
    "suporte tecnico",
    "sistemas internos",
    "atendimento ao cliente",
    "rotinas administrativas",
    "digitacao",
    "relatorios",
    "monitoramento",
    "escala",
    "call center",
    "power bi",
    "sql",
    "python",
]

CV_SECTION_HEADERS = {
    "experiencia",
    "experiencia profissional",
    "experiencias",
    "historico profissional",
    "atuacao profissional",
    "ultimas experiencias",
    "empregos anteriores",
    "empresa",
    "cargo",
    "periodo",
    "responsabilidades",
    "atividades",
}

NAME_LABEL_REGEX = re.compile(
    r"^(?:nome|nome completo|candidato|curriculo de)\s*[:\-]\s*(?P<nome>.+)$",
    re.IGNORECASE,
)


def normalize_cv_text(texto):
    bruto = str(texto or "")
    bruto = bruto.replace("\r", "\n")
    bruto = re.sub(r"\n{3,}", "\n\n", bruto)
    bruto = re.sub(r"[ \t]{2,}", " ", bruto)
    return bruto.strip()


def extract_email(texto):
    match = EMAIL_REGEX.search(texto or "")
    return match.group(1).strip() if match else ""


def extract_phone(texto):
    match = PHONE_REGEX.search(texto or "")
    return match.group(0).strip() if match else ""


def extract_whatsapp(texto):
    match = WHATSAPP_REGEX.search(texto or "")
    return match.group(0).strip() if match else ""


def is_valid_candidate_name(nome):
    nome_limpo = normalize_text(nome)
    if not nome_limpo:
        return False

    partes = [parte for parte in nome_limpo.split() if parte.strip()]
    if len(partes) < 2 or len(partes) > 5:
        return False

    if len(nome_limpo) < 8 or len(nome_limpo) > 60:
        return False

    termos_invalidos = {
        "curriculo",
        "objetivo",
        "resumo",
        "perfil",
        "experiencia",
        "formacao",
        "habilidades",
        "competencias",
        "telefone",
        "email",
        "e-mail",
        "whatsapp",
        "linkedin",
        "github",
        "endereco",
        "bairro",
        "cidade",
        "rua",
        "control desk",
        "atendente",
        "analista",
        "operador",
        "supervisor",
        "auxiliar",
        "assistente",
        "jovem aprendiz",
    }

    nome_normalizado = normalize_compare_text(nome_limpo)
    if any(termo in nome_normalizado for termo in termos_invalidos):
        return False

    if any(char.isdigit() for char in nome_limpo):
        return False
    return bool(re.fullmatch(r"[A-Za-zÀ-ÖØ-öø-ÿ' ]+", nome_limpo))


def _clean_candidate_name(value: str) -> str:
    cleaned = normalize_text(value)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    cleaned = cleaned.strip(" -.|:;")
    return cleaned


def is_valid_email(email):
    email_limpo = normalize_text(email)
    if not email_limpo:
        return False
    return bool(EMAIL_REGEX.fullmatch(email_limpo))


def sanitize_phone(valor):
    return re.sub(r"\D", "", str(valor or ""))


def is_valid_phone(telefone):
    numero = sanitize_phone(telefone)
    if numero.startswith("55") and len(numero) in (12, 13):
        numero = numero[2:]
    return len(numero) in (10, 11)


def extract_education_strength(texto):
    base = normalize_compare_text(texto)
    if any(item in base for item in ["mba", "pos-graduacao", "pos graduacao"]):
        return "forte"
    if any(item in base for item in ["graduacao", "bacharelado", "tecnologo", "faculdade", "universidade"]):
        return "media"
    if any(item in base for item in ["ensino medio", "curso tecnico", "escolaridade"]):
        return "basica"
    return "ausente"


def has_experience_content(texto):
    base = normalize_compare_text(texto)
    pistas = [
        "experiencia",
        "atuei",
        "trabalhei",
        "atuacao",
        "empresa",
        "cargo",
        "funcao",
        "responsavel por",
        "atividades desenvolvidas",
    ]
    return any(pista in base for pista in pistas)


def extract_experience_strength(texto):
    base = normalize_compare_text(texto)
    marcadores_tempo = re.findall(r"\b(19|20)\d{2}\b", base)
    tem_empresa = any(item in base for item in ["empresa", "cargo", "funcao", "responsavel por"])

    if len(marcadores_tempo) >= 2 and tem_empresa:
        return "forte"
    if len(marcadores_tempo) >= 1 or tem_empresa:
        return "media"
    if has_experience_content(texto):
        return "basica"
    return "ausente"


def extract_candidate_name_details(texto, fallback_name: str = "", filename: str = "") -> dict:
    safe_fallback = _clean_candidate_name(fallback_name)
    if is_valid_candidate_name(safe_fallback):
        return {"nome": safe_fallback, "confianca": "alta", "fonte": "formulario"}

    linhas = [linha.strip() for linha in str(texto or "").splitlines() if linha.strip()]
    candidatos = []

    for linha in linhas[:12]:
        label_match = NAME_LABEL_REGEX.search(linha)
        if label_match:
            linha_limpa = _clean_candidate_name(label_match.group("nome"))
            if is_valid_candidate_name(linha_limpa):
                candidatos.append((linha_limpa, 0, "rotulo"))

    for index, linha in enumerate(linhas[:10]):
        linha_limpa = _clean_candidate_name(linha)
        if is_valid_candidate_name(linha_limpa):
            candidatos.append((linha_limpa, index + 2, "primeiras_linhas"))

    if filename:
        stem = Path(normalize_text(filename)).stem
        stem = re.sub(r"(?i)\b(curriculo|cv|resume|candidato|vaga)\b", " ", stem)
        stem = re.sub(r"[-_.]+", " ", stem)
        stem = _clean_candidate_name(stem)
        if is_valid_candidate_name(stem):
            candidatos.append((stem, 25, "arquivo"))

    if not candidatos:
        return {"nome": "", "confianca": "baixa", "fonte": ""}

    candidatos.sort(key=lambda item: (item[1], len(item[0].split()), len(item[0])))
    nome, prioridade, fonte = candidatos[0]
    confianca = "alta" if prioridade <= 2 else "media"
    return {"nome": nome, "confianca": confianca, "fonte": fonte}


def extract_candidate_name(texto, fallback_name: str = "", filename: str = ""):
    return extract_candidate_name_details(texto, fallback_name, filename).get("nome", "")


def extract_competencies(texto):
    base = normalize_compare_text(texto)
    comportamentais = [item for item in BEHAVIORAL_COMPETENCIES if item in base]
    tecnicas = [item for item in TECHNICAL_COMPETENCIES if item in base]
    return {
        "comportamentais": sorted(set(comportamentais)),
        "tecnicas": sorted(set(tecnicas)),
    }


def extract_professional_experiences(texto):
    lines = [normalize_text(line) for line in str(texto or "").splitlines() if normalize_text(line)]
    base_lines = [normalize_compare_text(line) for line in lines]
    experiences = []
    in_section = False
    for index, line in enumerate(lines):
        normalized = base_lines[index]
        if normalized in CV_SECTION_HEADERS or any(header in normalized for header in CV_SECTION_HEADERS):
            in_section = True
            continue
        if in_section and normalized in {"formacao", "habilidades", "competencias", "cursos"}:
            break
        if in_section or any(marker in normalized for marker in ["empresa", "cargo", "periodo", "responsabilidades", "atividades"]):
            if len(line) >= 8 and len(experiences) < 8:
                experiences.append(line)
    return experiences


def extract_keywords(texto):
    base = normalize_compare_text(texto)
    encontrados = []
    termos_negativos = [
        "nao tenho experiencia com",
        "nao possuo experiencia com",
        "sem experiencia em",
        "nunca trabalhei com",
        "nao conheco",
        "nao domino",
    ]

    for palavra in CV_KEYWORDS:
        if palavra not in base:
            continue

        posicao = base.find(palavra)
        trecho_inicio = max(0, posicao - 60)
        trecho_fim = min(len(base), posicao + len(palavra) + 60)
        contexto = base[trecho_inicio:trecho_fim]

        if any(termo in contexto for termo in termos_negativos):
            continue

        encontrados.append(palavra)

    return sorted(set(encontrados))


def build_cv_justification(classificacao, keywords_validas, competencias, experiencias, problemas, pontos_fortes):
    partes = []
    if experiencias:
        partes.append("Currículo apresenta experiência profissional descrita, com atividades e histórico aproveitáveis para revisão do RH.")
    if keywords_validas:
        partes.append(f"Foram encontrados sinais aderentes a vaga: {', '.join(keywords_validas[:6])}.")
    comp = (competencias or {}).get("comportamentais", []) + (competencias or {}).get("tecnicas", [])
    if comp:
        partes.append(f"Competencias identificadas: {', '.join(comp[:8])}.")
    if problemas:
        partes.append(f"Pontos de atencao: {'; '.join(problemas[:3])}.")
    if normalize_compare_text(classificacao) == "nao qualificado":
        partes.append("Ainda assim, recomenda-se revisão manual do RH antes do descarte definitivo.")
    return " ".join(partes) or "Análise concluída com informações limitadas no currículo."


def score_cv_for_role(
    role,
    keywords_found,
    has_email,
    has_phone,
    text_length,
    candidate_name="",
    email_value="",
    phone_value="",
    education_strength="ausente",
    experience_strength="ausente",
    name_confidence="baixa",
    competencies=None,
    experiences=None,
):
    normalized_role = normalize_text(role)
    if normalize_compare_text(normalized_role) == "estagiario":
        normalized_role = "Estagiario"

    role_cfg = CV_WEIGHTS_BY_ROLE.get(
        normalized_role,
        {
            "keywords": {
                "atendimento": 1.5,
                "cliente": 1.2,
                "suporte": 1.2,
                "excel": 1.0,
                "word": 0.8,
                "crm": 1.0,
                "monitoramento": 1.0,
                "rotinas administrativas": 1.0,
                "comunicacao": 0.8,
                "organizacao": 0.8,
            },
            "min_keywords": 2,
        },
    )

    score = 0.0
    problemas = []
    pontos_fortes = []
    keyword_weights = role_cfg.get("keywords", {})
    keywords_validas = []

    for keyword in keywords_found:
        peso = float(keyword_weights.get(keyword, 0.0))
        if peso > 0:
            keywords_validas.append(keyword)
            score += peso

    if is_valid_candidate_name(candidate_name):
        score += 1.0 if name_confidence == "alta" else 0.75
        pontos_fortes.append(f"Nome do candidato identificado com confiança {name_confidence}.")
    else:
        score -= 2.0
        problemas.append("Nome do candidato não foi identificado com segurança.")

    if has_email and is_valid_email(email_value):
        score += 1.0
        pontos_fortes.append("E-mail válido identificado.")
    else:
        score -= 1.5
        problemas.append("E-mail não encontrado ou inválido.")

    if has_phone and is_valid_phone(phone_value):
        score += 1.0
        pontos_fortes.append("Telefone ou WhatsApp válido identificado.")
    else:
        score -= 1.5
        problemas.append("Telefone/WhatsApp não encontrado ou inválido.")

    if text_length < 80:
        score -= 3.0
        problemas.append("CV com pouco conteúdo extraído.")
    elif text_length < 180:
        score -= 1.5
        problemas.append("CV com conteúdo muito limitado para análise segura.")
    elif text_length >= 250:
        score += 0.5

    min_keywords = int(role_cfg.get("min_keywords", 2))
    if len(keywords_validas) >= min_keywords:
        score += 0.5
        pontos_fortes.append(
            f"Foram identificadas {len(keywords_validas)} palavra(s)-chave aderentes a vaga."
        )
    else:
        score -= 2.0
        problemas.append("Poucas palavras-chave realmente aderentes a vaga.")

    if not keywords_validas:
        score -= 2.0
        problemas.append("Nenhuma palavra-chave relevante da vaga foi validada no currículo.")

    if education_strength == "forte":
        score += 1.5
        pontos_fortes.append("Formação acadêmica robusta identificada.")
    elif education_strength == "media":
        score += 0.9
        pontos_fortes.append("Formação acadêmica compatível identificada.")
    elif education_strength == "basica":
        score += 0.3
        pontos_fortes.append("Indícios básicos de formação identificados.")
    else:
        score -= 1.2
        problemas.append("Formação/educação não identificada no currículo.")

    if experience_strength == "forte":
        score += 2.0
        pontos_fortes.append("Experiência profissional consistente identificada.")
    elif experience_strength == "media":
        score += 1.0
        pontos_fortes.append("Experiência profissional parcial identificada.")
    elif experience_strength == "basica":
        score += 0.3
        pontos_fortes.append("Há menções superficiais de experiência profissional.")
    else:
        score -= 1.8
        problemas.append("Experiência profissional não identificada com clareza.")

    competencias = competencies or {"comportamentais": [], "tecnicas": []}
    total_competencias = len(competencias.get("comportamentais", [])) + len(competencias.get("tecnicas", []))
    if total_competencias >= 4:
        score += 1.2
        pontos_fortes.append("Conjunto consistente de competências comportamentais e técnicas identificado.")
    elif total_competencias >= 2:
        score += 0.6
        pontos_fortes.append("Algumas competências relevantes foram identificadas.")
    else:
        problemas.append("Poucas competências explícitas foram encontradas no currículo.")

    experiences = experiences or []
    if len(experiences) >= 3:
        score += 1.0
        pontos_fortes.append("Experiencias profissionais com bom nivel de detalhe foram identificadas.")
    elif len(experiences) >= 1:
        score += 0.4

    score = max(0.0, min(10.0, round(score, 2)))

    if (
        score >= 7.0
        and len(keywords_validas) >= min_keywords
        and is_valid_candidate_name(candidate_name)
        and is_valid_email(email_value)
        and is_valid_phone(phone_value)
        and experience_strength in ("forte", "media")
    ):
        classificacao = "Qualificado"
        slug = "qualificado"
    elif (
        score >= 4.5
        and len(keywords_validas) >= max(1, min_keywords - 1)
        and experience_strength in ("forte", "media", "basica")
    ):
        classificacao = "Parcialmente qualificado"
        slug = "parcialmente-qualificado"
    else:
        classificacao = "Não qualificado"
        slug = "nao-qualificado"

    return {
        "score": score,
        "classificacao": classificacao,
        "slug": slug,
        "problemas": problemas,
        "pontos_fortes": pontos_fortes,
        "keywords_validas": keywords_validas,
        "education_strength": education_strength,
        "experience_strength": experience_strength,
        "competencias": competencias,
        "experiencias": experiences,
        "justificativa": build_cv_justification(
            classificacao,
            keywords_validas,
            competencias,
            experiences,
            problemas,
            pontos_fortes,
        ),
    }


def _detect_cv_file(filename: str, content_type: str = "") -> tuple[str, str]:
    safe_name = normalize_text(filename)
    extension = Path(safe_name).suffix.lower()
    mime_type = normalize_text(content_type).lower()
    if not mime_type:
        guessed_type, _ = mimetypes.guess_type(safe_name)
        mime_type = normalize_text(guessed_type).lower()
    return extension, mime_type or "application/octet-stream"


def _ensure_supported_cv_file(filename: str, content_bytes: bytes, content_type: str = "") -> tuple[str, str, bytes]:
    extension, mime_type = _detect_cv_file(filename, content_type)
    if extension not in SUPPORTED_CV_EXTENSIONS:
        raise CvTextExtractionError(
            "Formato de currículo não suportado. Envie um arquivo PDF, DOC ou DOCX."
        )

    safe_content = content_bytes or b""
    if not safe_content:
        raise CvTextExtractionError("Não foi possível encontrar texto no currículo enviado.")

    if len(safe_content) > CV_MAX_BYTES:
        raise CvTextExtractionError("Arquivo excede o tamanho máximo permitido para análise de CV.")

    return extension, mime_type, safe_content


def _normalize_extracted_text(text: str) -> str:
    return normalize_cv_text(text)


def _extract_txt(content_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise CvTextExtractionError("Arquivo TXT vazio ou corrompido.", technical_message="TXT decoding failed.")


def _extract_pdf(content_bytes: bytes, filename: str) -> str:
    try:
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError as exc:
                raise CvTextExtractionError(
                    "Biblioteca pypdf/PyPDF2 não instalada.",
                    technical_message="Neither pypdf nor PyPDF2 is installed.",
                ) from exc

        reader = PdfReader(io.BytesIO(content_bytes))
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")
            except Exception as exc:
                raise CvTextExtractionError(
                    "Não foi possível ler o currículo. Verifique se o arquivo não está corrompido e tente novamente.",
                    technical_message=f"PDF is encrypted and could not be decrypted: {exc}",
                ) from exc
        parts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
        text = _normalize_extracted_text("\n".join(parts))
        if text:
            return text
    except CvTextExtractionError:
        raise
    except PermissionError as exc:
        raise CvTextExtractionError(
            "Falha ao ler o arquivo por permissao de acesso.",
            technical_message=f"Permission denied reading PDF: {exc}",
        ) from exc
    except Exception as exc:
        raise CvTextExtractionError(
            "Não foi possível ler o currículo. Verifique se o arquivo não está corrompido e tente novamente.",
            technical_message=f"PDF extraction failed: {type(exc).__name__}: {exc}",
        ) from exc

    raise CvTextExtractionError(
        "Não foi possível extrair texto deste PDF. Verifique se o arquivo possui texto selecionável e não é apenas uma imagem escaneada.",
        technical_message="PDF has no selectable text.",
    )


def _extract_docx(content_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise CvTextExtractionError(
            "Biblioteca python-docx não instalada.",
            technical_message="python-docx is not installed.",
        ) from exc

    try:
        document = Document(io.BytesIO(content_bytes))
    except (BadZipFile, ValueError, KeyError) as exc:
        raise CvTextExtractionError(
            "Não foi possível ler o currículo. Verifique se o arquivo não está corrompido e tente novamente.",
            technical_message=f"DOCX is invalid: {type(exc).__name__}: {exc}",
        ) from exc
    except PermissionError as exc:
        raise CvTextExtractionError(
            "Falha ao ler o arquivo por permissao de acesso.",
            technical_message=f"Permission denied reading DOCX: {exc}",
        ) from exc
    except Exception as exc:
        raise CvTextExtractionError(
            "Não foi possível ler o currículo. Verifique se o arquivo não está corrompido e tente novamente.",
            technical_message=f"DOCX extraction failed: {type(exc).__name__}: {exc}",
        ) from exc

    parts = [paragraph.text for paragraph in document.paragraphs if normalize_text(paragraph.text)]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = normalize_text(cell.text)
                if cell_text:
                    parts.append(cell_text)
    return _normalize_extracted_text("\n".join(parts))


def _extract_doc_with_command(command: list[str], content_bytes: bytes, suffix: str) -> str:
    with tempfile.TemporaryDirectory(prefix="rh-cv-doc-") as temp_dir:
        input_path = Path(temp_dir) / f"curriculo{suffix}"
        input_path.write_bytes(content_bytes)
        completed = subprocess.run(
            command(input_path) if callable(command) else command,
            capture_output=True,
            text=True,
            timeout=30,
            check=False,
        )
        if completed.returncode != 0:
            logger.info(
                "Conversor DOC retornou codigo %s: %s",
                completed.returncode,
                normalize_text(completed.stderr),
            )
            return ""
        return _normalize_extracted_text(completed.stdout)


def _extract_doc_with_libreoffice(content_bytes: bytes) -> str:
    _, libreoffice_path = _get_doc_converter_settings()
    soffice = (
        normalize_text(libreoffice_path)
        or shutil.which("soffice")
        or shutil.which("libreoffice")
    )
    if not soffice:
        return ""

    with tempfile.TemporaryDirectory(prefix="rh-cv-doc-lo-") as temp_dir:
        input_path = Path(temp_dir) / "curriculo.doc"
        output_dir = Path(temp_dir) / "out"
        output_dir.mkdir(parents=True, exist_ok=True)
        input_path.write_bytes(content_bytes)
        completed = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(output_dir),
                str(input_path),
            ],
            capture_output=True,
            text=True,
            timeout=45,
            check=False,
        )
        if completed.returncode != 0:
            logger.info("LibreOffice DOC conversion failed: %s", normalize_text(completed.stderr))
            return ""
        output_files = list(output_dir.glob("*.txt"))
        if not output_files:
            return ""
        return _extract_txt(output_files[0].read_bytes())


def _extract_doc_with_word_com(content_bytes: bytes) -> str:
    try:
        import win32com.client  # type: ignore[import-not-found]
    except ImportError:
        return ""

    word = None
    document = None
    with tempfile.TemporaryDirectory(prefix="rh-cv-doc-word-") as temp_dir:
        input_path = Path(temp_dir) / "curriculo.doc"
        output_path = Path(temp_dir) / "curriculo.txt"
        input_path.write_bytes(content_bytes)
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(
                str(input_path),
                ReadOnly=True,
                AddToRecentFiles=False,
                ConfirmConversions=False,
            )
            try:
                if hasattr(document, "SaveAs2"):
                    document.SaveAs2(str(output_path), FileFormat=2)
                else:
                    document.SaveAs(str(output_path), FileFormat=2)
            finally:
                if document is not None:
                    document.Close(False)
            if output_path.exists():
                return _extract_txt(output_path.read_bytes())
            return ""
        except Exception as exc:
            logger.info("Microsoft Word COM DOC extraction failed: %s", exc)
            return ""
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    logger.info("Falha ao encerrar Microsoft Word apos extracao de DOC.")


def _get_doc_converter_settings() -> tuple[str, str]:
    converter = normalize_compare_text(os.getenv("DOC_CONVERTER") or os.getenv("RH_CV_DOC_CONVERTER") or "")
    libreoffice_path = normalize_text(os.getenv("LIBREOFFICE_PATH") or os.getenv("RH_LIBREOFFICE_PATH") or "")

    if not converter or not libreoffice_path:
        try:
            from ..config import get_settings

            settings = get_settings()
            converter = converter or normalize_compare_text(settings.doc_converter)
            libreoffice_path = libreoffice_path or normalize_text(settings.libreoffice_path)
        except Exception as exc:
            logger.info("Não foi possível ler a configuração do conversor DOC: %s", exc)

    if converter not in {"auto", "word", "libreoffice", "disabled"}:
        logger.info("DOC_CONVERTER inválido (%s); usando modo auto.", converter)
        converter = "auto"

    return converter or "auto", libreoffice_path


def _extract_doc(content_bytes: bytes) -> str:
    converter, _ = _get_doc_converter_settings()
    if converter == "disabled":
        raise CvTextExtractionError(
            DOC_CONVERTER_FAILURE_MESSAGE,
            technical_message="DOC converter disabled by configuration.",
        )

    if converter in {"auto", "word"}:
        text = _extract_doc_with_word_com(content_bytes)
        if text:
            return text

    if converter in {"auto", "libreoffice"}:
        text = _extract_doc_with_libreoffice(content_bytes)
        if text:
            return text

    raise CvTextExtractionError(
        DOC_CONVERTER_FAILURE_MESSAGE,
        technical_message=(
            "No DOC converter available. Configure Microsoft Word COM or LibreOffice "
            "with DOC_CONVERTER=auto|word|libreoffice."
        ),
    )


def _extract_rtf(content_bytes: bytes) -> str:
    raw_text = _extract_txt(content_bytes)
    try:
        from striprtf.striprtf import rtf_to_text

        return _normalize_extracted_text(rtf_to_text(raw_text))
    except ImportError:
        logger.info("striprtf não instalado; usando fallback simples para RTF.")
    except Exception as exc:
        logger.info("Falha no striprtf para RTF: %s", exc)

    without_controls = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw_text)
    without_controls = re.sub(r"\\[a-zA-Z]+\d* ?", " ", without_controls)
    without_controls = without_controls.replace("{", " ").replace("}", " ")
    return _normalize_extracted_text(without_controls)


def _extract_odt(content_bytes: bytes) -> str:
    try:
        from odf import text as odf_text
        from odf.opendocument import load
    except ImportError as exc:
        raise CvTextExtractionError(
            "Biblioteca odfpy não instalada.",
            technical_message="odfpy is not installed.",
        ) from exc

    try:
        document = load(io.BytesIO(content_bytes))
        parts = []
        for element in document.getElementsByType(odf_text.P):
            parts.append("".join(node.data for node in element.childNodes if hasattr(node, "data")))
        for element in document.getElementsByType(odf_text.H):
            parts.append("".join(node.data for node in element.childNodes if hasattr(node, "data")))
        return _normalize_extracted_text("\n".join(parts))
    except PermissionError as exc:
        raise CvTextExtractionError(
            "Falha ao ler o arquivo por permissao de acesso.",
            technical_message=f"Permission denied reading ODT: {exc}",
        ) from exc
    except Exception as exc:
        raise CvTextExtractionError(
            "Arquivo vazio ou corrompido.",
            technical_message=f"ODT extraction failed: {type(exc).__name__}: {exc}",
        ) from exc


def _try_ocr_image_or_pdf(content_bytes: bytes, filename: str, extension: str) -> str:
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return ""

    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.info(
            "OCR indisponivel para %s: instale pytesseract e Pillow; para PDF escaneado, instale tambem pdf2image/poppler.",
            filename,
        )
        return ""

    if extension == ".pdf":
        try:
            from pdf2image import convert_from_bytes
        except ImportError:
            logger.info("OCR de PDF indisponível para %s: pdf2image/poppler não instalado.", filename)
            return ""
        try:
            images = convert_from_bytes(content_bytes, first_page=1, last_page=5)
            return _normalize_extracted_text("\n".join(pytesseract.image_to_string(image) for image in images))
        except Exception as exc:
            logger.info("OCR de PDF falhou para %s: %s", filename, exc)
            return ""

    try:
        image = Image.open(io.BytesIO(content_bytes))
        return _normalize_extracted_text(pytesseract.image_to_string(image))
    except Exception as exc:
        logger.info("OCR de imagem falhou para %s: %s", filename, exc)
        return ""


def extract_text_from_uploaded_file(filename, content_bytes, content_type=""):
    extension, mime_type, safe_content = _ensure_supported_cv_file(filename, content_bytes, content_type)
    logger.info(
        "Iniciando extracao de CV: arquivo=%s extensao=%s mime=%s tamanho_bytes=%s",
        normalize_text(filename) or "(sem nome)",
        extension,
        mime_type,
        len(safe_content),
    )

    try:
        if extension == ".txt":
            text = _extract_txt(safe_content)
        elif extension == ".pdf":
            text = _extract_pdf(safe_content, normalize_text(filename))
        elif extension == ".docx":
            text = _extract_docx(safe_content)
        elif extension == ".doc":
            text = _extract_doc(safe_content)
        elif extension == ".rtf":
            text = _extract_rtf(safe_content)
        elif extension == ".odt":
            text = _extract_odt(safe_content)
        elif extension in IMAGE_CV_EXTENSIONS:
            text = _try_ocr_image_or_pdf(safe_content, normalize_text(filename), extension)
            if not text:
                raise CvTextExtractionError("Imagem recebida, mas OCR não está habilitado neste servidor.")
        else:
            raise CvTextExtractionError(
                "Formato de currículo não suportado. Envie um arquivo PDF, DOC ou DOCX."
            )
    except CvTextExtractionError:
        raise
    except PermissionError as exc:
        raise CvTextExtractionError(
            "Falha ao ler o arquivo por permissao de acesso.",
            technical_message=f"Permission denied extracting CV: {exc}",
        ) from exc
    except Exception as exc:
        raise CvTextExtractionError(
            "Não foi possível ler o currículo. Verifique se o arquivo não está corrompido e tente novamente.",
            technical_message=f"Unexpected CV extraction error: {type(exc).__name__}: {exc}",
        ) from exc

    text = _normalize_extracted_text(text)
    if not text:
        raise CvTextExtractionError(
            "Não foi possível encontrar texto no currículo enviado.",
            technical_message="Extractor returned empty text.",
        )

    logger.info(
        "Extracao de CV concluida: arquivo=%s extensao=%s caracteres=%s",
        normalize_text(filename) or "(sem nome)",
        extension,
        len(text),
    )
    return text


def serialize_cv_problems(avaliacao: dict) -> str:
    return json.dumps(
        {
            "problemas": avaliacao["problemas"],
            "pontos_fortes": avaliacao["pontos_fortes"],
            "education_strength": avaliacao["education_strength"],
            "experience_strength": avaliacao["experience_strength"],
            "competencias": avaliacao.get("competencias", {}),
            "experiencias": avaliacao.get("experiencias", []),
            "justificativa": avaliacao.get("justificativa", ""),
            "nome_detectado": avaliacao.get("nome_detectado", ""),
            "confianca_nome": avaliacao.get("confianca_nome", ""),
            "fonte_nome": avaliacao.get("fonte_nome", ""),
        },
        ensure_ascii=False,
    )
