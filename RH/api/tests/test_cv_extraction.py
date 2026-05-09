from __future__ import annotations

import unittest
import io
import sys
from pathlib import Path
from unittest.mock import patch

API_DIR = Path(__file__).resolve().parents[1]
if str(API_DIR) not in sys.path:
    sys.path.insert(0, str(API_DIR))

from rh_api.services.cv import CvTextExtractionError, extract_text_from_uploaded_file


class CvExtractionTests(unittest.TestCase):
    def test_extract_docx(self):
        from docx import Document

        document = Document()
        document.add_paragraph("Joao Silva")
        document.add_paragraph("Experiencia com atendimento")
        buffer = io.BytesIO()
        document.save(buffer)

        text = extract_text_from_uploaded_file(
            "curriculo.docx",
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        self.assertIn("Joao Silva", text)
        self.assertIn("atendimento", text)

    def test_rejects_empty_file_with_specific_message(self):
        with self.assertRaises(CvTextExtractionError) as context:
            extract_text_from_uploaded_file("curriculo.docx", b"", "")

        self.assertEqual(
            context.exception.user_message,
            "Não foi possível encontrar texto no currículo enviado.",
        )

    def test_rejects_unknown_extension_with_specific_message(self):
        for filename in ("curriculo.png", "curriculo.txt", "curriculo.xlsx"):
            with self.assertRaises(CvTextExtractionError) as context:
                extract_text_from_uploaded_file(filename, b"conteudo", "application/octet-stream")

            self.assertEqual(
                context.exception.user_message,
                "Formato de currículo não suportado. Envie um arquivo PDF, DOC ou DOCX.",
            )

    def test_pdf_without_selectable_text_returns_specific_message(self):
        from pypdf import PdfWriter

        buffer = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        writer.write(buffer)

        with self.assertRaises(CvTextExtractionError) as context:
            extract_text_from_uploaded_file("curriculo.pdf", buffer.getvalue(), "application/pdf")

        self.assertEqual(
            context.exception.user_message,
            "Não foi possível extrair texto deste PDF. Verifique se o arquivo possui texto selecionável e não é apenas uma imagem escaneada.",
        )

    def test_doc_without_converter_returns_specific_message(self):
        with patch("rh_api.services.cv._extract_doc_with_word_com", return_value=""):
            with patch("rh_api.services.cv.shutil.which", return_value=None):
                with self.assertRaises(CvTextExtractionError) as context:
                    extract_text_from_uploaded_file(
                        "curriculo.doc",
                        b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1conteudo",
                        "application/msword",
                    )

        self.assertEqual(
            context.exception.user_message,
            "Não foi possível converter o arquivo .doc neste servidor. Verifique se o Microsoft Word ou LibreOffice está instalado/configurado, ou envie o currículo em PDF ou DOCX.",
        )


if __name__ == "__main__":
    unittest.main()
