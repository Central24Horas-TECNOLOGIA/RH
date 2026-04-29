from __future__ import annotations

import io
import json
import logging
import mimetypes
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from zipfile import BadZipFile

from .helpers import normalize_compare_text, normalize_text


logger = logging.getLogger(__name__)

CV_MAX_BYTES = 10 * 1024 * 1024
SUPPORTED_CV_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".rtf",
    ".odt",
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".bmp",
}
IMAGE_CV_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}


class CvTextExtractionError(Exception):
    def __init__(self, user_message: str, *, technical_message: str = ""):
        super().__init__(technical_message or user_message)
        self.user_message = user_message
        self.technical_message = technical_message or user_message


CV_KEYWORDS = [
    "excel",
    "word",
    "power bi",
    "powerbi",
    "atendimento",
    "suporte",
    "planilha",
    "dashboards",
    "dashboard",
    "sql",
    "python",
    "javascript",
    "react",
    "talkdesk",
    "power automate",
    "sharepoint",
    "teams",
    "analise",
    "dados",
    "cliente",
    "monitoramento",
]

CV_WEIGHTS_BY_ROLE = {
    "Jovem Aprendiz": {
        "keywords": {
            "excel": 1.5,
            "word": 1.0,
            "atendimento": 2.0,
            "cliente": 1.5,
            "dados": 1.0,
        },
        "min_keywords": 2,
    },
    "Operador": {
        "keywords": {
            "atendimento": 2.5,
            "cliente": 2.0,
            "suporte": 1.5,
            "talkdesk": 2.0,
            "monitoramento": 1.0,
            "excel": 1.0,
        },
        "min_keywords": 2,
    },
    "Supervisor": {
        "keywords": {
            "excel": 2.0,
            "atendimento": 2.0,
            "cliente": 1.5,
            "monitoramento": 1.5,
            "dados": 1.5,
            "power bi": 2.0,
        },
        "min_keywords": 3,
    },
    "Estagiario": {
        "keywords": {
            "excel": 1.5,
            "python": 2.0,
            "sql": 2.0,
            "dados": 1.5,
            "javascript": 1.5,
            "react": 1.5,
        },
        "min_keywords": 2,
    },
    "Analista": {
        "keywords": {
            "excel": 1.5,
            "power bi": 2.5,
            "sql": 2.0,
            "python": 2.0,
            "dados": 2.0,
            "analise": 1.5,
        },
        "min_keywords": 3,
    },
    "TI": {
        "keywords": {
            "python": 2.5,
            "javascript": 2.0,
            "react": 2.0,
            "sql": 2.0,
            "sharepoint": 1.5,
            "power automate": 1.5,
            "teams": 1.0,
            "suporte": 1.5,
        },
        "min_keywords": 3,
    },
}

EMAIL_REGEX = re.compile(r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})", re.IGNORECASE)
PHONE_REGEX = re.compile(r"(?:(?:\+?55)\s*)?(?:\(?\d{2}\)?\s*)?(?:9?\d{4})[-\s.]?\d{4}")
WHATSAPP_REGEX = re.compile(r"(?:(?:\+?55)\s*)?(?:\(?\d{2}\)?\s*)?(?:9\d{4})[-\s.]?\d{4}")


def normalize_cv_text(texto):
    bruto = str(texto or "")
    bruto = bruto.replace("\r", "\n")
    bruto = re.sub(r"\n{3,}", "\n\n", bruto)
    bruto = re.sub(r"[ \t]{2,}", " ", bruto)
    return bruto.strip()


def extract_email(texto):
    match = EMAIL_REGEX.search(texto or "")
    return match.group(1).strip() if match else ""


def extract_phone(texto):
    match = PHONE_REGEX.search(texto or "")
    return match.group(0).strip() if match else ""


def extract_whatsapp(texto):
    match = WHATSAPP_REGEX.search(texto or "")
    return match.group(0).strip() if match else ""


def is_valid_candidate_name(nome):
    nome_limpo = normalize_text(nome)
    if not nome_limpo:
        return False

    partes = [parte for parte in nome_limpo.split() if parte.strip()]
    if len(partes) < 2 or len(partes) > 5:
        return False

    if len(nome_limpo) < 8 or len(nome_limpo) > 60:
        return False

    termos_invalidos = {
        "curriculo",
        "objetivo",
        "resumo",
        "perfil",
        "experiencia",
        "formacao",
        "habilidades",
        "competencias",
        "telefone",
        "email",
        "e-mail",
        "whatsapp",
        "linkedin",
        "github",
    }

    nome_normalizado = normalize_compare_text(nome_limpo)
    if any(termo in nome_normalizado for termo in termos_invalidos):
        return False

    return not any(char.isdigit() for char in nome_limpo)


def is_valid_email(email):
    email_limpo = normalize_text(email)
    if not email_limpo:
        return False
    return bool(EMAIL_REGEX.fullmatch(email_limpo))


def sanitize_phone(valor):
    return re.sub(r"\D", "", str(valor or ""))


def is_valid_phone(telefone):
    numero = sanitize_phone(telefone)
    if numero.startswith("55") and len(numero) in (12, 13):
        numero = numero[2:]
    return len(numero) in (10, 11)


def extract_education_strength(texto):
    base = normalize_compare_text(texto)
    if any(item in base for item in ["mba", "pos-graduacao", "pos graduacao"]):
        return "forte"
    if any(item in base for item in ["graduacao", "bacharelado", "tecnologo", "faculdade", "universidade"]):
        return "media"
    if any(item in base for item in ["ensino medio", "curso tecnico", "escolaridade"]):
        return "basica"
    return "ausente"


def has_experience_content(texto):
    base = normalize_compare_text(texto)
    pistas = [
        "experiencia",
        "atuei",
        "trabalhei",
        "atuacao",
        "empresa",
        "cargo",
        "funcao",
        "responsavel por",
        "atividades desenvolvidas",
    ]
    return any(pista in base for pista in pistas)


def extract_experience_strength(texto):
    base = normalize_compare_text(texto)
    marcadores_tempo = re.findall(r"\b(19|20)\d{2}\b", base)
    tem_empresa = any(item in base for item in ["empresa", "cargo", "funcao", "responsavel por"])

    if len(marcadores_tempo) >= 2 and tem_empresa:
        return "forte"
    if len(marcadores_tempo) >= 1 or tem_empresa:
        return "media"
    if has_experience_content(texto):
        return "basica"
    return "ausente"


def extract_candidate_name(texto):
    linhas = [linha.strip() for linha in str(texto or "").splitlines() if linha.strip()]
    if not linhas:
        return ""

    candidatos = []
    for linha in linhas[:15]:
        linha_limpa = linha.strip(" -.|")
        if is_valid_candidate_name(linha_limpa):
            candidatos.append(linha_limpa)

    if not candidatos:
        return ""

    candidatos.sort(key=lambda item: (len(item.split()), len(item)))
    return candidatos[0]


def extract_keywords(texto):
    base = normalize_compare_text(texto)
    encontrados = []
    termos_negativos = [
        "nao tenho experiencia com",
        "nao possuo experiencia com",
        "sem experiencia em",
        "nunca trabalhei com",
        "nao conheco",
        "nao domino",
    ]

    for palavra in CV_KEYWORDS:
        if palavra not in base:
            continue

        posicao = base.find(palavra)
        trecho_inicio = max(0, posicao - 60)
        trecho_fim = min(len(base), posicao + len(palavra) + 60)
        contexto = base[trecho_inicio:trecho_fim]

        if any(termo in contexto for termo in termos_negativos):
            continue

        encontrados.append(palavra)

    return sorted(set(encontrados))


def score_cv_for_role(
    role,
    keywords_found,
    has_email,
    has_phone,
    text_length,
    candidate_name="",
    email_value="",
    phone_value="",
    education_strength="ausente",
    experience_strength="ausente",
):
    normalized_role = normalize_text(role)
    if normalize_compare_text(normalized_role) == "estagiario":
        normalized_role = "Estagiario"

    role_cfg = CV_WEIGHTS_BY_ROLE.get(
        normalized_role,
        {"keywords": {}, "min_keywords": 2},
    )

    score = 0.0
    problemas = []
    pontos_fortes = []
    keyword_weights = role_cfg.get("keywords", {})
    keywords_validas = []

    for keyword in keywords_found:
        peso = float(keyword_weights.get(keyword, 0.0))
        if peso > 0:
            keywords_validas.append(keyword)
            score += peso

    if is_valid_candidate_name(candidate_name):
        score += 0.75
        pontos_fortes.append("Nome do candidato identificado com consistencia.")
    else:
        score -= 2.0
        problemas.append("Nome do candidato nao foi identificado com seguranca.")

    if has_email and is_valid_email(email_value):
        score += 1.0
        pontos_fortes.append("E-mail valido identificado.")
    else:
        score -= 1.5
        problemas.append("E-mail nao encontrado ou invalido.")

    if has_phone and is_valid_phone(phone_value):
        score += 1.0
        pontos_fortes.append("Telefone ou WhatsApp valido identificado.")
    else:
        score -= 1.5
        problemas.append("Telefone/WhatsApp nao encontrado ou invalido.")

    if text_length < 80:
        score -= 3.0
        problemas.append("CV com pouco conteudo extraido.")
    elif text_length < 180:
        score -= 1.5
        problemas.append("CV com conteudo muito limitado para analise segura.")
    elif text_length >= 250:
        score += 0.5

    min_keywords = int(role_cfg.get("min_keywords", 2))
    if len(keywords_validas) >= min_keywords:
        score += 0.5
        pontos_fortes.append(
            f"Foram identificadas {len(keywords_validas)} palavra(s)-chave aderentes a vaga."
        )
    else:
        score -= 2.0
        problemas.append("Poucas palavras-chave realmente aderentes a vaga.")

    if not keywords_validas:
        score -= 2.0
        problemas.append("Nenhuma palavra-chave relevante da vaga foi validada no curriculo.")

    if education_strength == "forte":
        score += 1.5
        pontos_fortes.append("Formacao academica robusta identificada.")
    elif education_strength == "media":
        score += 0.9
        pontos_fortes.append("Formacao academica compativel identificada.")
    elif education_strength == "basica":
        score += 0.3
        pontos_fortes.append("Indicios basicos de formacao identificados.")
    else:
        score -= 1.2
        problemas.append("Formacao/educacao nao identificada no curriculo.")

    if experience_strength == "forte":
        score += 2.0
        pontos_fortes.append("Experiencia profissional consistente identificada.")
    elif experience_strength == "media":
        score += 1.0
        pontos_fortes.append("Experiencia profissional parcial identificada.")
    elif experience_strength == "basica":
        score += 0.3
        pontos_fortes.append("Ha mencoes superficiais de experiencia profissional.")
    else:
        score -= 1.8
        problemas.append("Experiencia profissional nao identificada com clareza.")

    score = max(0.0, min(10.0, round(score, 2)))

    if (
        score >= 7.0
        and len(keywords_validas) >= min_keywords
        and is_valid_candidate_name(candidate_name)
        and is_valid_email(email_value)
        and is_valid_phone(phone_value)
        and experience_strength in ("forte", "media")
    ):
        classificacao = "Qualificado"
        slug = "qualificado"
    elif (
        score >= 4.5
        and len(keywords_validas) >= max(1, min_keywords - 1)
        and experience_strength in ("forte", "media", "basica")
    ):
        classificacao = "Qualificado"
        slug = "qualificado"
    else:
        classificacao = "Nao qualificado"
        slug = "nao-qualificado"

    return {
        "score": score,
        "classificacao": classificacao,
        "slug": slug,
        "problemas": problemas,
        "pontos_fortes": pontos_fortes,
        "keywords_validas": keywords_validas,
        "education_strength": education_strength,
        "experience_strength": experience_strength,
    }


def _detect_cv_file(filename: str, content_type: str = "") -> tuple[str, str]:
    safe_name = normalize_text(filename)
    extension = Path(safe_name).suffix.lower()
    mime_type = normalize_text(content_type).lower()
    if not mime_type:
        guessed_type, _ = mimetypes.guess_type(safe_name)
        mime_type = normalize_text(guessed_type).lower()
    return extension, mime_type or "application/octet-stream"


def _ensure_supported_cv_file(filename: str, content_bytes: bytes, content_type: str = "") -> tuple[str, str, bytes]:
    extension, mime_type = _detect_cv_file(filename, content_type)
    if extension not in SUPPORTED_CV_EXTENSIONS:
        raise CvTextExtractionError(
            "Formato nao suportado para analise automatica. Envie PDF com texto selecionavel, DOCX, DOC, TXT, RTF ou ODT."
        )

    safe_content = content_bytes or b""
    if not safe_content:
        raise CvTextExtractionError("Arquivo vazio ou corrompido.")

    if len(safe_content) > CV_MAX_BYTES:
        raise CvTextExtractionError("Arquivo excede o tamanho maximo permitido para analise de CV.")

    return extension, mime_type, safe_content


def _normalize_extracted_text(text: str) -> str:
    return normalize_cv_text(text)


def _extract_txt(content_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise CvTextExtractionError("Arquivo TXT vazio ou corrompido.", technical_message="TXT decoding failed.")


def _extract_pdf(content_bytes: bytes, filename: str) -> str:
    try:
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError as exc:
                raise CvTextExtractionError(
                    "Biblioteca pypdf/PyPDF2 nao instalada.",
                    technical_message="Neither pypdf nor PyPDF2 is installed.",
                ) from exc

        reader = PdfReader(io.BytesIO(content_bytes))
        parts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
        text = _normalize_extracted_text("\n".join(parts))
        if text:
            return text
    except CvTextExtractionError:
        raise
    except PermissionError as exc:
        raise CvTextExtractionError(
            "Falha ao ler o arquivo por permissao de acesso.",
            technical_message=f"Permission denied reading PDF: {exc}",
        ) from exc
    except Exception as exc:
        raise CvTextExtractionError(
            "Arquivo vazio ou corrompido.",
            technical_message=f"PDF extraction failed: {type(exc).__name__}: {exc}",
        ) from exc

    ocr_text = _try_ocr_image_or_pdf(content_bytes, filename, ".pdf")
    if ocr_text:
        return ocr_text

    raise CvTextExtractionError("PDF sem texto selecionavel. OCR nao esta habilitado neste servidor.")


def _extract_docx(content_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise CvTextExtractionError(
            "Biblioteca python-docx nao instalada.",
            technical_message="python-docx is not installed.",
        ) from exc

    try:
        document = Document(io.BytesIO(content_bytes))
    except (BadZipFile, ValueError, KeyError) as exc:
        raise CvTextExtractionError(
            "Arquivo vazio ou corrompido.",
            technical_message=f"DOCX is invalid: {type(exc).__name__}: {exc}",
        ) from exc
    except PermissionError as exc:
        raise CvTextExtractionError(
            "Falha ao ler o arquivo por permissao de acesso.",
            technical_message=f"Permission denied reading DOCX: {exc}",
        ) from exc
    except Exception as exc:
        raise CvTextExtractionError(
            "Arquivo vazio ou corrompido.",
            technical_message=f"DOCX extraction failed: {type(exc).__name__}: {exc}",
        ) from exc

    parts = [paragraph.text for paragraph in document.paragraphs if normalize_text(paragraph.text)]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = normalize_text(cell.text)
                if cell_text:
                    parts.append(cell_text)
    return _normalize_extracted_text("\n".join(parts))


def _extract_doc_with_command(command: list[str], content_bytes: bytes, suffix: str) -> str:
    with tempfile.TemporaryDirectory(prefix="rh-cv-doc-") as temp_dir:
        input_path = Path(temp_dir) / f"curriculo{suffix}"
        input_path.write_bytes(content_bytes)
        completed = subprocess.run(
            command(input_path) if callable(command) else command,
            capture_output=True,
            text=True,
            timeout=30,
            check=False,
        )
        if completed.returncode != 0:
            logger.info(
                "Conversor DOC retornou codigo %s: %s",
                completed.returncode,
                normalize_text(completed.stderr),
            )
            return ""
        return _normalize_extracted_text(completed.stdout)


def _extract_doc_with_libreoffice(content_bytes: bytes) -> str:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return ""

    with tempfile.TemporaryDirectory(prefix="rh-cv-doc-lo-") as temp_dir:
        input_path = Path(temp_dir) / "curriculo.doc"
        output_dir = Path(temp_dir) / "out"
        output_dir.mkdir(parents=True, exist_ok=True)
        input_path.write_bytes(content_bytes)
        completed = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(output_dir),
                str(input_path),
            ],
            capture_output=True,
            text=True,
            timeout=45,
            check=False,
        )
        if completed.returncode != 0:
            logger.info("LibreOffice DOC conversion failed: %s", normalize_text(completed.stderr))
            return ""
        output_files = list(output_dir.glob("*.txt"))
        if not output_files:
            return ""
        return _extract_txt(output_files[0].read_bytes())


def _extract_doc(content_bytes: bytes) -> str:
    antiword = shutil.which("antiword")
    if antiword:
        text = _extract_doc_with_command(lambda path: [antiword, str(path)], content_bytes, ".doc")
        if text:
            return text

    text = _extract_doc_with_libreoffice(content_bytes)
    if text:
        return text

    raise CvTextExtractionError(
        "Arquivo .doc antigo nao pode ser lido neste ambiente. Converta para .docx ou habilite o conversor definido no backend.",
        technical_message="No DOC extractor available. Install antiword or LibreOffice and expose it in PATH.",
    )


def _extract_rtf(content_bytes: bytes) -> str:
    raw_text = _extract_txt(content_bytes)
    try:
        from striprtf.striprtf import rtf_to_text

        return _normalize_extracted_text(rtf_to_text(raw_text))
    except ImportError:
        logger.info("striprtf nao instalado; usando fallback simples para RTF.")
    except Exception as exc:
        logger.info("Falha no striprtf para RTF: %s", exc)

    without_controls = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw_text)
    without_controls = re.sub(r"\\[a-zA-Z]+\d* ?", " ", without_controls)
    without_controls = without_controls.replace("{", " ").replace("}", " ")
    return _normalize_extracted_text(without_controls)


def _extract_odt(content_bytes: bytes) -> str:
    try:
        from odf import text as odf_text
        from odf.opendocument import load
    except ImportError as exc:
        raise CvTextExtractionError(
            "Biblioteca odfpy nao instalada.",
            technical_message="odfpy is not installed.",
        ) from exc

    try:
        document = load(io.BytesIO(content_bytes))
        parts = []
        for element in document.getElementsByType(odf_text.P):
            parts.append("".join(node.data for node in element.childNodes if hasattr(node, "data")))
        for element in document.getElementsByType(odf_text.H):
            parts.append("".join(node.data for node in element.childNodes if hasattr(node, "data")))
        return _normalize_extracted_text("\n".join(parts))
    except PermissionError as exc:
        raise CvTextExtractionError(
            "Falha ao ler o arquivo por permissao de acesso.",
            technical_message=f"Permission denied reading ODT: {exc}",
        ) from exc
    except Exception as exc:
        raise CvTextExtractionError(
            "Arquivo vazio ou corrompido.",
            technical_message=f"ODT extraction failed: {type(exc).__name__}: {exc}",
        ) from exc


def _try_ocr_image_or_pdf(content_bytes: bytes, filename: str, extension: str) -> str:
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return ""

    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.info(
            "OCR indisponivel para %s: instale pytesseract e Pillow; para PDF escaneado, instale tambem pdf2image/poppler.",
            filename,
        )
        return ""

    if extension == ".pdf":
        try:
            from pdf2image import convert_from_bytes
        except ImportError:
            logger.info("OCR de PDF indisponivel para %s: pdf2image/poppler nao instalado.", filename)
            return ""
        try:
            images = convert_from_bytes(content_bytes, first_page=1, last_page=5)
            return _normalize_extracted_text("\n".join(pytesseract.image_to_string(image) for image in images))
        except Exception as exc:
            logger.info("OCR de PDF falhou para %s: %s", filename, exc)
            return ""

    try:
        image = Image.open(io.BytesIO(content_bytes))
        return _normalize_extracted_text(pytesseract.image_to_string(image))
    except Exception as exc:
        logger.info("OCR de imagem falhou para %s: %s", filename, exc)
        return ""


def extract_text_from_uploaded_file(filename, content_bytes, content_type=""):
    extension, mime_type, safe_content = _ensure_supported_cv_file(filename, content_bytes, content_type)
    logger.info(
        "Iniciando extracao de CV: arquivo=%s extensao=%s mime=%s tamanho_bytes=%s",
        normalize_text(filename) or "(sem nome)",
        extension,
        mime_type,
        len(safe_content),
    )

    try:
        if extension == ".txt":
            text = _extract_txt(safe_content)
        elif extension == ".pdf":
            text = _extract_pdf(safe_content, normalize_text(filename))
        elif extension == ".docx":
            text = _extract_docx(safe_content)
        elif extension == ".doc":
            text = _extract_doc(safe_content)
        elif extension == ".rtf":
            text = _extract_rtf(safe_content)
        elif extension == ".odt":
            text = _extract_odt(safe_content)
        elif extension in IMAGE_CV_EXTENSIONS:
            text = _try_ocr_image_or_pdf(safe_content, normalize_text(filename), extension)
            if not text:
                raise CvTextExtractionError("Imagem recebida, mas OCR nao esta habilitado neste servidor.")
        else:
            raise CvTextExtractionError(
                "Formato nao suportado para analise automatica. Envie PDF com texto selecionavel, DOCX, DOC, TXT, RTF ou ODT."
            )
    except CvTextExtractionError:
        raise
    except PermissionError as exc:
        raise CvTextExtractionError(
            "Falha ao ler o arquivo por permissao de acesso.",
            technical_message=f"Permission denied extracting CV: {exc}",
        ) from exc
    except Exception as exc:
        raise CvTextExtractionError(
            "Arquivo vazio ou corrompido.",
            technical_message=f"Unexpected CV extraction error: {type(exc).__name__}: {exc}",
        ) from exc

    text = _normalize_extracted_text(text)
    if not text:
        raise CvTextExtractionError("Arquivo vazio ou corrompido.", technical_message="Extractor returned empty text.")

    logger.info(
        "Extracao de CV concluida: arquivo=%s extensao=%s caracteres=%s",
        normalize_text(filename) or "(sem nome)",
        extension,
        len(text),
    )
    return text


def serialize_cv_problems(avaliacao: dict) -> str:
    return json.dumps(
        {
            "problemas": avaliacao["problemas"],
            "pontos_fortes": avaliacao["pontos_fortes"],
            "education_strength": avaliacao["education_strength"],
            "experience_strength": avaliacao["experience_strength"],
        },
        ensure_ascii=False,
    )
